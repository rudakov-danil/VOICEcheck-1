"""
REST API routes for dialog export functionality.
"""

import logging
from typing import Optional
from pathlib import Path
from datetime import datetime
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from ..database import models as db_models
from ..database.connection import get_db

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/export", tags=["export"])


@router.get("/dialogs/{dialog_id}/pdf")
async def export_dialog_pdf(
    dialog_id: str,
    db: AsyncSession = Depends(get_db)
):
    """
    Export dialog analysis to PDF format.

    Args:
        dialog_id: Dialog ID
        db: Database session

    Returns:
        FileResponse: PDF file for download
    """
    try:
        # Get dialog with analysis
        dialog_uuid = UUID(dialog_id)
        query = select(db_models.Dialog).where(
            db_models.Dialog.id == dialog_uuid
        ).options(
            selectinload(db_models.Dialog.transcriptions),
            selectinload(db_models.Dialog.analyses)
        )

        result = await db.execute(query)
        dialog = result.scalar_one_or_none()

        if not dialog:
            raise HTTPException(status_code=404, detail="Dialog not found")

        if not dialog.analyses:
            raise HTTPException(status_code=400, detail="Analysis not available for this dialog")

        # Generate PDF
        pdf_file = await generate_pdf_export(dialog)

        return FileResponse(
            path=pdf_file,
            filename=f"{dialog.filename}_analysis.pdf",
            media_type="application/pdf"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to export PDF for dialog {dialog_id}: {e}")
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")


@router.get("/dialogs/{dialog_id}/docx")
async def export_dialog_docx(
    dialog_id: str,
    db: AsyncSession = Depends(get_db)
):
    """
    Export dialog analysis to DOCX format.

    Args:
        dialog_id: Dialog ID
        db: Database session

    Returns:
        FileResponse: DOCX file for download
    """
    try:
        # Get dialog with analysis
        dialog_uuid = UUID(dialog_id)
        query = select(db_models.Dialog).where(
            db_models.Dialog.id == dialog_uuid
        ).options(
            selectinload(db_models.Dialog.transcriptions),
            selectinload(db_models.Dialog.analyses)
        )

        result = await db.execute(query)
        dialog = result.scalar_one_or_none()

        if not dialog:
            raise HTTPException(status_code=404, detail="Dialog not found")

        if not dialog.analyses:
            raise HTTPException(status_code=400, detail="Analysis not available for this dialog")

        # Generate DOCX
        docx_file = await generate_docx_export(dialog)

        return FileResponse(
            path=docx_file,
            filename=f"{dialog.filename}_analysis.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to export DOCX for dialog {dialog_id}: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")


async def generate_pdf_export(dialog: db_models.Dialog) -> str:
    """
    Generate PDF export file for dialog analysis.

    Args:
        dialog: Dialog model with analysis

    Returns:
        str: Path to generated PDF file
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Create temporary directory for exports
        export_dir = Path("/app/exports")
        export_dir.mkdir(parents=True, exist_ok=True)

        # Generate unique filename
        filename = f"{dialog.id}_analysis.pdf"
        pdf_file = export_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(str(pdf_file), pagesize=A4, rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)

        # Create styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.darkblue
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )

        # Add Russian font support
        try:
            pdfmetrics.registerFont(TTFont('Arial', '/usr/share/fonts/truetype/msttcorefonts/Arial.ttf'))
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                fontName='Arial',
                spaceAfter=6
            )
        except:
            body_style = styles['Normal']

        # Build content
        content = []

        # Title
        title = Paragraph(f"VOICEcheck - Анализ диалога", title_style)
        content.append(title)
        content.append(Spacer(1, 12))

        # Dialog info
        content.append(Paragraph(f"Файл: {dialog.filename}", heading_style))
        content.append(Paragraph(f"Дата: {dialog.created_at.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
        content.append(Paragraph(f"Длительность: {dialog.duration:.1f} секунд", body_style))
        content.append(Spacer(1, 12))

        if dialog.analyses:
            analysis = dialog.analyses[0]

            # Scores section
            content.append(Paragraph("Оценки по категориям:", heading_style))
            scores_data = [['Категория', 'Оценка (0-10)']]
            for category, score in analysis.scores.items():
                scores_data.append([category.capitalize(), f"{score:.1f}"])

            scores_table = Table(scores_data, colWidths=[3*inch, 1.5*inch])
            scores_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            content.append(scores_table)
            content.append(Spacer(1, 12))

            # Key moments section
            content.append(Paragraph("Ключевые моменты:", heading_style))
            for moment in analysis.key_moments:
                time_text = f"[{moment['time']:.0f}s]"
                moment_text = f"{time_text} {moment.get('text', '')}"
                content.append(Paragraph(moment_text, body_style))
            content.append(Spacer(1, 12))

            # Recommendations section
            content.append(Paragraph("Рекомендации:", heading_style))
            for i, rec in enumerate(analysis.recommendations, 1):
                content.append(Paragraph(f"{i}. {rec}", body_style))
            content.append(Spacer(1, 12))

            # Speaking time section
            content.append(Paragraph("Статистика времени речи:", heading_style))
            speaking_time = analysis.speaking_time
            content.append(Paragraph(f"Продавец: {speaking_time.get('sales', 0):.1f} секунд", body_style))
            content.append(Paragraph(f"Клиент: {speaking_time.get('customer', 0):.1f} секунд", body_style))

        # Build PDF
        doc.build(content)

        logger.info(f"PDF generated successfully: {pdf_file}")
        return str(pdf_file)

    except ImportError:
        # Fallback to simple text file if reportlab not available
        return await generate_text_export(dialog, "pdf")
    except Exception as e:
        logger.error(f"Error generating PDF: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate PDF")


async def generate_docx_export(dialog: db_models.Dialog) -> str:
    """
    Generate DOCX export file for dialog analysis.

    Args:
        dialog: Dialog model with analysis

    Returns:
        str: Path to generated DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create temporary directory for exports
        export_dir = Path("/app/exports")
        export_dir.mkdir(parents=True, exist_ok=True)

        # Generate unique filename
        filename = f"{dialog.id}_analysis.docx"
        docx_file = export_dir / filename

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading('VOICEcheck - Анализ диалога', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Dialog info
        info_para = doc.add_paragraph()
        info_para.add_run(f"Файл: {dialog.filename}\n").bold = True
        info_para.add_run(f"Дата: {dialog.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Длительность: {dialog.duration:.1f} секунд\n")

        if dialog.analyses:
            analysis = dialog.analyses[0]

            # Scores section
            scores_heading = doc.add_heading('Оценки по категориям:', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Категория'
            hdr_cells[1].text = 'Оценка (0-10)'

            # Data
            for category, score in analysis.scores.items():
                row_cells = table.add_row().cells
                row_cells[0].text = category.capitalize()
                row_cells[1].text = f"{score:.1f}"

            # Key moments section
            moments_heading = doc.add_heading('Ключевые моменты:', level=1)
            for moment in analysis.key_moments:
                time_text = f"[{moment['time']:.0f}s]"
                moment_para = doc.add_paragraph()
                moment_para.add_run(time_text).bold = True
                moment_para.add_run(f" {moment.get('text', '')}")

            # Recommendations section
            recs_heading = doc.add_heading('Рекомендации:', level=1)
            for i, rec in enumerate(analysis.recommendations, 1):
                doc.add_paragraph(f"{i}. {rec}", style='List Number')

            # Speaking time section
            speaking_heading = doc.add_heading('Статистика времени речи:', level=1)
            speaking_para = doc.add_paragraph()
            speaking_para.add_run("Продавец: ").bold = True
            speaking_para.add_run(f"{analysis.speaking_time.get('sales', 0):.1f} секунд\n")
            speaking_para.add_run("Клиент: ").bold = True
            speaking_para.add_run(f"{analysis.speaking_time.get('customer', 0):.1f} секунд")

        # Save document
        doc.save(str(docx_file))

        logger.info(f"DOCX generated successfully: {docx_file}")
        return str(docx_file)

    except ImportError:
        # Fallback to simple text file if python-docx not available
        return await generate_text_export(dialog, "docx")
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX")


async def generate_text_export(dialog: db_models.Dialog, format: str) -> str:
    """
    Generate simple text file as fallback for export.

    Args:
        dialog: Dialog model with analysis
        format: Export format for filename

    Returns:
        str: Path to generated text file
    """
    from pathlib import Path
    import tempfile

    # Create temporary directory for exports
    export_dir = Path("/app/exports")
    export_dir.mkdir(parents=True, exist_ok=True)

    # Generate unique filename
    filename = f"{dialog.id}_analysis.{format}"
    export_file = export_dir / filename

    # Generate text content
    with open(export_file, "w", encoding="utf-8") as f:
        f.write(f"VOICEcheck - Анализ диалога\n")
        f.write(f"Файл: {dialog.filename}\n")
        f.write(f"Дата: {dialog.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Длительность: {dialog.duration} секунд\n\n")

        if dialog.analyses:
            analysis = dialog.analyses[0]
            f.write("Оценки по категориям:\n")
            for category, score in analysis.scores.items():
                f.write(f"  {category}: {score}/10\n")

            f.write("\nКлючевые моменты:\n")
            for moment in analysis.key_moments:
                f.write(f"  [{moment['time']:.0f}s] {moment['text']}\n")

            f.write("\nРекомендации:\n")
            for i, rec in enumerate(analysis.recommendations, 1):
                f.write(f"  {i}. {rec}\n")

            f.write("\nСтатистика времени речи:\n")
            f.write(f"  Продавец: {analysis.speaking_time.get('sales', 0)} секунд\n")
            f.write(f"  Клиент: {analysis.speaking_time.get('customer', 0)} секунд\n")

    return str(export_file)